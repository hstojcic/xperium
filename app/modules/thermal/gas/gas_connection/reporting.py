"""
Funkcije za izvještavanje i izvoz podataka.
"""
import math
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from .constants import *

def export_to_word(doc, data):
    """
    Izvozi proračun u Word dokument.
    
    Args:
        doc: python-docx Document objekt
        data: Dictionary s podacima proračuna
    """
    # Stiliziranje naslova
    style = doc.styles['Title']
    style.font.size = Pt(16)
    style.font.bold = True
    
    # Naslov dokumenta
    doc.add_heading("Proračun plinskog priključka", 0)
    
    # Datum i oznaka dokumenta
    from datetime import datetime
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Datum: {datetime.now().strftime('%d.%m.%Y.')}")
    p.add_run("\nOznaka: PP-" + datetime.now().strftime('%Y%m%d-%H%M'))
    
    # Intro
    doc.add_paragraph("Ovaj dokument sadrži proračun plinskog priključka prema važećim tehničkim propisima za plinske instalacije.")
    
    # 1. Osnovni podaci
    doc.add_heading("1. Osnovni podaci", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Vrsta plinskog priključka: ").bold = True
    p.add_run("Niskotlačni (do 100 mbar)" if data["tip_prikljucka"] == 1 else "Srednjetlačni (0,1-5,0 bar)")
    
    p = doc.add_paragraph()
    p.add_run("Duljina cjevovoda: ").bold = True
    p.add_run(f"{data['cijev']['duljina']:.2f} m")
    
    # 2. Podaci o uređajima
    doc.add_heading("2. Odabrani uređaji", level=1)
    
    # Podaci o kotlu
    if data["kotao"]["ima_kotao"]:
        doc.add_heading("2.1. Plinski kotao", level=2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Formatiranje tablice
        for cell in table.rows[0].cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Zaglavlje
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Parametar"
        hdr_cells[1].text = "Vrijednost"
        
        # Podaci
        row_cells = table.add_row().cells
        row_cells[0].text = "Tip kotla"
        row_cells[1].text = "Kondenzacijski" if data["kotao"]["tip"] == 1 else "Klasični"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Proizvođač"
        row_cells[1].text = data["kotao"]["proizvodjac"]
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Model"
        row_cells[1].text = data["kotao"]["model"]
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Broj kotlova"
        row_cells[1].text = str(data["kotao"]["broj_kotlova"])
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Snaga grijanja po kotlu"
        row_cells[1].text = f"{data['kotao']['snaga_grijanja']:.2f} kW"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Snaga PTV-a po kotlu"
        row_cells[1].text = f"{data['kotao']['snaga_ptv']:.2f} kW"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Ukupna snaga grijanja"
        row_cells[1].text = f"{data['kotao']['snaga_grijanja'] * data['kotao']['broj_kotlova']:.2f} kW"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Ukupna snaga PTV-a"
        row_cells[1].text = f"{data['kotao']['snaga_ptv'] * data['kotao']['broj_kotlova']:.2f} kW"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Faktor istovremenosti"
        row_cells[1].text = f"{data['kotao']['faktor']:.3f}"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Iskoristivost"
        row_cells[1].text = f"{data['kotao']['eta']:.2f}"
    
    # Podaci o plinskom uređaju
    if data["plinski_uredjaj"]["ima_uredjaj"]:
        doc.add_heading(f"2.2. {data['plinski_uredjaj']['naziv']}", level=2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Formatiranje tablice
        for cell in table.rows[0].cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Zaglavlje
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Parametar"
        hdr_cells[1].text = "Vrijednost"
        
        # Podaci
        row_cells = table.add_row().cells
        row_cells[0].text = "Broj uređaja"
        row_cells[1].text = str(data["plinski_uredjaj"]["broj_uredjaja"])
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Broj plamenika"
        row_cells[1].text = str(data["plinski_uredjaj"]["broj_plamenika"])
        
        if data["plinski_uredjaj"]["ima_pecnicu"]:
            row_cells = table.add_row().cells
            row_cells[0].text = "Pećnica"
            row_cells[1].text = "Da"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Snaga po uređaju"
        row_cells[1].text = f"{data['plinski_uredjaj']['snaga_jedinicna']:.2f} kW"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Ukupna snaga"
        row_cells[1].text = f"{data['plinski_uredjaj']['snaga_ukupna']:.2f} kW"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Faktor istovremenosti"
        row_cells[1].text = f"{data['plinski_uredjaj']['faktor']:.3f}"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Iskoristivost"
        row_cells[1].text = f"{data['plinski_uredjaj']['eta']:.2f}"
    
    # 3. Proračun
    doc.add_heading("3. Proračun", level=1)
    
    # 3.1 Proračun vršnog protoka
    doc.add_heading("3.1. Proračun vršnog protoka", level=2)
    
    doc.add_paragraph("Formula za izračun vršnog protoka: Q = (P × f) / (Hd × η)")
    
    doc.add_paragraph(f"Gdje je:")
    p = doc.add_paragraph()
    p.add_run("Q").italic = True
    p.add_run(" - vršni protok plina [m³/h]")
    p = doc.add_paragraph()
    p.add_run("P").italic = True
    p.add_run(" - snaga uređaja [kW]")
    p = doc.add_paragraph()
    p.add_run("f").italic = True
    p.add_run(" - faktor istovremenosti")
    p = doc.add_paragraph()
    p.add_run("Hd").italic = True
    p.add_run(f" - ogrijevna vrijednost plina = {HD} kWh/m³")
    p = doc.add_paragraph()
    p.add_run("η").italic = True
    p.add_run(" - iskoristivost uređaja")
    
    # Proračun za kotao
    if data["kotao"]["ima_kotao"]:
        P_koristeno = max(
            data["kotao"]["snaga_grijanja"],
            data["kotao"]["snaga_ptv"]
        ) * data["kotao"]["broj_kotlova"]
        
        doc.add_paragraph("Proračun za plinski kotao:")
        
        p = doc.add_paragraph()
        p.add_run(f"P = {P_koristeno:.2f} kW")
        p = doc.add_paragraph()
        p.add_run(f"f = {data['kotao']['faktor']:.3f}")
        p = doc.add_paragraph()
        p.add_run(f"Hd = {HD:.2f} kWh/m³")
        p = doc.add_paragraph()
        p.add_run(f"η = {data['kotao']['eta']:.2f}")
        
        vrsni_protok_kotao = (P_koristeno * data["kotao"]["faktor"]) / (HD * data["kotao"]["eta"])
        
        p = doc.add_paragraph()
        p.add_run(f"Qk = (P × f) / (Hd × η) = ({P_koristeno:.2f} × {data['kotao']['faktor']:.3f}) / ({HD:.2f} × {data['kotao']['eta']:.2f}) = {vrsni_protok_kotao:.2f} m³/h")
    
    # Proračun za plinski uređaj
    if data["plinski_uredjaj"]["ima_uredjaj"]:
        doc.add_paragraph(f"Proračun za {data['plinski_uredjaj']['naziv'].lower()}:")
        
        p = doc.add_paragraph()
        p.add_run(f"P = {data['plinski_uredjaj']['snaga_ukupna']:.2f} kW")
        p = doc.add_paragraph()
        p.add_run(f"f = {data['plinski_uredjaj']['faktor']:.3f}")
        p = doc.add_paragraph()
        p.add_run(f"Hd = {HD:.2f} kWh/m³")
        p = doc.add_paragraph()
        p.add_run(f"η = {data['plinski_uredjaj']['eta']:.2f}")
        
        vrsni_protok_uredjaj = (data["plinski_uredjaj"]["snaga_ukupna"] * data["plinski_uredjaj"]["faktor"]) / (HD * data["plinski_uredjaj"]["eta"])
        
        p = doc.add_paragraph()
        p.add_run(f"Qu = (P × f) / (Hd × η) = ({data['plinski_uredjaj']['snaga_ukupna']:.2f} × {data['plinski_uredjaj']['faktor']:.3f}) / ({HD:.2f} × {data['plinski_uredjaj']['eta']:.2f}) = {vrsni_protok_uredjaj:.2f} m³/h")
    
    # Ukupni vršni protok
    doc.add_paragraph("Ukupni vršni protok:")
    p = doc.add_paragraph()
    p.add_run(f"Q = {data['rezultati']['vrsni_protok']:.2f} m³/h").bold = True
    
    # 3.2 Proračun potrebnog promjera cijevi
    doc.add_heading("3.2. Proračun potrebnog promjera cijevi", level=2)
    
    doc.add_paragraph("Formula za izračun potrebnog promjera cijevi: d = √(4 × Q)/(π × w)")
    
    doc.add_paragraph(f"Gdje je:")
    p = doc.add_paragraph()
    p.add_run("d").italic = True
    p.add_run(" - potreban promjer cijevi [mm]")
    p = doc.add_paragraph()
    p.add_run("Q").italic = True
    p.add_run(" - vršni protok plina [m³/s]")
    p = doc.add_paragraph()
    p.add_run("w").italic = True
    p.add_run(f" - preporučena brzina plina = {STANDARD_BRZINA} m/s")
    
    # Proračun potrebnog promjera
    p = doc.add_paragraph()
    p.add_run(f"Q = {data['rezultati']['vrsni_protok']:.2f} m³/h = {data['rezultati']['vrsni_protok']/3600:.6f} m³/s")
    p = doc.add_paragraph()
    p.add_run(f"w = {STANDARD_BRZINA:.1f} m/s")
    p = doc.add_paragraph()
    p.add_run(f"d = √(4 × {data['rezultati']['vrsni_protok']/3600:.6f})/(π × {STANDARD_BRZINA:.1f}) = {data['rezultati']['potreban_promjer']:.2f} mm").bold = True
    
    # 3.3 Odabir cijevi
    if data["cijev"]["oznaka"]:
        doc.add_heading("3.3. Odabir cijevi", level=2)
        
        naziv_cijevi = "PE-HD" if data["cijev"]["tip"] == 1 else "SMLS"
        
        p = doc.add_paragraph()
        p.add_run(f"Odabrana cijev: {naziv_cijevi} {data['cijev']['oznaka']}")
        p = doc.add_paragraph()
        p.add_run(f"Dimenzije: Ø{data['cijev']['dimenzije']['vanjski']}×{data['cijev']['dimenzije']['debljina']} mm")
        p = doc.add_paragraph()
        p.add_run(f"Unutarnji promjer: {data['cijev']['dimenzije']['unutarnji']:.1f} mm")
    
        # 3.4 Proračun stvarne brzine
        doc.add_heading("3.4. Proračun stvarne brzine u cijevi", level=2)
        
        doc.add_paragraph("Formula za izračun stvarne brzine u cijevi: w = (4 × Q)/(π × d²)")
        
        doc.add_paragraph(f"Gdje je:")
        p = doc.add_paragraph()
        p.add_run("w").italic = True
        p.add_run(" - stvarna brzina plina [m/s]")
        p = doc.add_paragraph()
        p.add_run("Q").italic = True
        p.add_run(" - vršni protok plina [m³/s]")
        p = doc.add_paragraph()
        p.add_run("d").italic = True
        p.add_run(" - unutarnji promjer cijevi [m]")
        
        # Proračun stvarne brzine
        p = doc.add_paragraph()
        p.add_run(f"Q = {data['rezultati']['vrsni_protok']:.2f} m³/h = {data['rezultati']['vrsni_protok']/3600:.6f} m³/s")
        p = doc.add_paragraph()
        p.add_run(f"d = {data['cijev']['dimenzije']['unutarnji']:.2f} mm = {data['cijev']['dimenzije']['unutarnji']/1000:.6f} m")
        p = doc.add_paragraph()
        p.add_run(f"w = (4 × {data['rezultati']['vrsni_protok']/3600:.6f})/(π × {data['cijev']['dimenzije']['unutarnji']/1000:.6f}²) = {data['rezultati']['stvarna_brzina']:.2f} m/s").bold = True
        
        # 3.5 Proračun pada tlaka
        doc.add_heading("3.5. Proračun pada tlaka", level=2)
        
        lambda_koef = LAMBDA_PEHD if data["cijev"]["tip"] == 1 else LAMBDA_SMLS
        naziv_cijevi = "PE-HD" if data["cijev"]["tip"] == 1 else "SMLS"
        
        # Različite formule ovisno o tipu priključka
        if data["tip_prikljucka"] == 1:  # Niskotlačni
            doc.add_paragraph("Formula za izračun pada tlaka (niskotlačni priključak): Δp = 6.25 × λ × (Q²ᵥ × ρₚₗ × L) / (100 × dᵤ)⁵")
            
            doc.add_paragraph(f"Gdje je:")
            p = doc.add_paragraph()
            p.add_run("Δp").italic = True
            p.add_run(" - pad tlaka [mbar]")
            p = doc.add_paragraph()
            p.add_run("λ").italic = True
            p.add_run(f" - koeficijent trenja = {lambda_koef}")
            p = doc.add_paragraph()
            p.add_run("Qᵥ").italic = True
            p.add_run(" - vršni protok plina [m³/h]")
            p = doc.add_paragraph()
            p.add_run("ρₚₗ").italic = True
            p.add_run(f" - gustoća plina = {RHO_PLIN} kg/m³")
            p = doc.add_paragraph()
            p.add_run("L").italic = True
            p.add_run(" - duljina cjevovoda [m]")
            p = doc.add_paragraph()
            p.add_run("dᵤ").italic = True
            p.add_run(" - unutarnji promjer cijevi [m]")
            
            # Proračun pada tlaka
            p = doc.add_paragraph()
            p.add_run(f"λ = {lambda_koef}")
            p = doc.add_paragraph()
            p.add_run(f"Qᵥ = {data['rezultati']['vrsni_protok']:.2f} m³/h")
            p = doc.add_paragraph()
            p.add_run(f"ρₚₗ = {RHO_PLIN} kg/m³")
            p = doc.add_paragraph()
            p.add_run(f"L = {data['cijev']['duljina']:.2f} m")
            p = doc.add_paragraph()
            p.add_run(f"dᵤ = {data['cijev']['dimenzije']['unutarnji']:.2f} mm = {data['cijev']['dimenzije']['unutarnji']/1000:.6f} m")
            p = doc.add_paragraph()
            p.add_run(f"Δp = 6.25 × {lambda_koef} × ({data['rezultati']['vrsni_protok']:.2f}² × {RHO_PLIN} × {data['cijev']['duljina']:.2f}) / (100 × {data['cijev']['dimenzije']['unutarnji']/1000:.6f})⁵ = {data['rezultati']['pad_tlaka']:.4f} mbar").bold = True
            
        else:  # Srednjetlačni
            doc.add_paragraph("Formula za izračun pada tlaka (srednjetlačni priključak): p₁² - p₂² = (λ × L × w₁² × ρₚₗ,₁ × p₁) / dᵤ")
            
            doc.add_paragraph(f"Gdje je:")
            p = doc.add_paragraph()
            p.add_run("p₁").italic = True
            p.add_run(f" - početni tlak = {POCETNI_TLAK_SREDNJETLACNI} mbar")
            p = doc.add_paragraph()
            p.add_run("p₂").italic = True
            p.add_run(" - krajnji tlak [mbar]")
            p = doc.add_paragraph()
            p.add_run("λ").italic = True
            p.add_run(f" - koeficijent trenja = {lambda_koef}")
            p = doc.add_paragraph()
            p.add_run("L").italic = True
            p.add_run(" - duljina cjevovoda [m]")
            p = doc.add_paragraph()
            p.add_run("w₁").italic = True
            p.add_run(" - brzina plina [m/s]")
            p = doc.add_paragraph()
            p.add_run("ρₚₗ,₁").italic = True
            p.add_run(f" - gustoća plina = {RHO_PLIN} kg/m³")
            p = doc.add_paragraph()
            p.add_run("dᵤ").italic = True
            p.add_run(" - unutarnji promjer cijevi [m]")
            
            # Proračun pada tlaka
            p = doc.add_paragraph()
            p.add_run(f"λ = {lambda_koef}")
            p = doc.add_paragraph()
            p.add_run(f"L = {data['cijev']['duljina']:.2f} m")
            p = doc.add_paragraph()
            p.add_run(f"w₁ = {data['rezultati']['stvarna_brzina']:.2f} m/s")
            p = doc.add_paragraph()
            p.add_run(f"ρₚₗ,₁ = {RHO_PLIN} kg/m³")
            p = doc.add_paragraph()
            p.add_run(f"dᵤ = {data['cijev']['dimenzije']['unutarnji']:.2f} mm = {data['cijev']['dimenzije']['unutarnji']/1000:.6f} m")
            p = doc.add_paragraph()
            p.add_run(f"p₁ = {POCETNI_TLAK_SREDNJETLACNI} mbar")
            
            desna_strana = (lambda_koef * data['cijev']['duljina'] * pow(data['rezultati']['stvarna_brzina'], 2) * RHO_PLIN * POCETNI_TLAK_SREDNJETLACNI) / (data['cijev']['dimenzije']['unutarnji']/1000)
            p2_kvadrat = pow(POCETNI_TLAK_SREDNJETLACNI, 2) - desna_strana
            p2 = math.sqrt(p2_kvadrat)
            
            p = doc.add_paragraph()
            p.add_run(f"p₂ = {p2:.2f} mbar")
            p = doc.add_paragraph()
            p.add_run(f"Pad tlaka: Δp = p₁ - p₂ = {POCETNI_TLAK_SREDNJETLACNI} - {p2:.2f} = {data['rezultati']['pad_tlaka']:.2f} mbar").bold = True
    
    # 4. Odabir plinomjera
    if data["rezultati"]["plinomjer"]:
        doc.add_heading("4. Odabir plinomjera", level=1)
        
        oznaka, spec = data["rezultati"]["plinomjer"]
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Formatiranje tablice
        for cell in table.rows[0].cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Zaglavlje
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Parametar"
        hdr_cells[1].text = "Vrijednost"
        
        # Podaci
        row_cells = table.add_row().cells
        row_cells[0].text = "Oznaka plinomjera"
        row_cells[1].text = oznaka
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Nazivni promjer"
        row_cells[1].text = f"DN {spec['DN']}"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Minimalni protok"
        row_cells[1].text = f"{spec['Qmin']} m³/h"
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Maksimalni protok"
        row_cells[1].text = f"{spec['Qmax']} m³/h"
        
        # Provjera odgovara li plinomjer potrebnom protoku
        if data["rezultati"]["vrsni_protok"] > spec['Qmax']:
            p = doc.add_paragraph()
            run = p.add_run("UPOZORENJE: Vršni protok premašuje maksimalni protok plinomjera!")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    # 5. Zaključak
    doc.add_heading("5. Zaključak", level=1)
    
    if data["cijev"]["oznaka"]:
        max_pad = MAX_PAD_TLAKA_NISKOTLACNI if data["tip_prikljucka"] == 1 else MAX_PAD_TLAKA_SREDNJETLACNI
        je_u_granicama = data["rezultati"]["pad_tlaka"] <= max_pad
        
        p = doc.add_paragraph()
        if je_u_granicama:
            run = p.add_run("Proračun zadovoljava maksimalni dozvoljeni pad tlaka.")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)  # Zelena
        else:
            run = p.add_run("Proračun NE zadovoljava maksimalni dozvoljeni pad tlaka! Potrebno je odabrati veći promjer cijevi.")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Crvena
        
        p = doc.add_paragraph()
        p.add_run(f"Vršni protok: {data['rezultati']['vrsni_protok']:.2f} m³/h")
        p = doc.add_paragraph()
        p.add_run(f"Potreban promjer: {data['rezultati']['potreban_promjer']:.2f} mm")
        p = doc.add_paragraph()
        naziv_cijevi = "PE-HD" if data["cijev"]["tip"] == 1 else "SMLS"
        p.add_run(f"Odabrana cijev: {naziv_cijevi} {data['cijev']['oznaka']} (unutarnji promjer: {data['cijev']['dimenzije']['unutarnji']:.1f} mm)")
        p = doc.add_paragraph()
        p.add_run(f"Stvarna brzina: {data['rezultati']['stvarna_brzina']:.2f} m/s")
        p = doc.add_paragraph()
        p.add_run(f"Pad tlaka: {data['rezultati']['pad_tlaka']:.2f} mbar")
        p = doc.add_paragraph()
        p.add_run(f"Dozvoljeni pad tlaka: {max_pad:.1f} mbar")